"""Create the Task 9 MER methods and interpretation Word note.

This is a lightweight documentation builder. It reads the Task 9 CSV outputs and
existing MER inventories; it does not run analysis workflows.
"""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("D:/Github_repos/Gayini")
MER_REPORT_DIR = ROOT / "Output" / "reports" / "MER"
MER_CSV_DIR = ROOT / "Output" / "csv" / "MER"
MER_FIG_DIR = ROOT / "Output" / "figures" / "review" / "MER"
OUT_DOCX = MER_REPORT_DIR / "Gayini_MER_methods_and_interpretation_note.docx"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text or "")
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
      set_cell_text(hdr[i], header, bold=True)
      set_cell_shading(hdr[i], "E8EEF5")
      if widths:
          hdr[i].width = Inches(widths[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])

    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    if path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(6.2))
        caption_p = doc.add_paragraph(caption)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.runs[0].italic = True
        caption_p.runs[0].font.size = Pt(9)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def main() -> None:
    MER_REPORT_DIR.mkdir(parents=True, exist_ok=True)

    agreement = read_csv(MER_CSV_DIR / "mer_vs_annual_occurrence_agreement_summary.csv")
    metric_defs = read_csv(MER_CSV_DIR / "mer_metric_definitions.csv")
    input_files = read_csv(MER_CSV_DIR / "mer_input_files_inventory.csv")
    output_files = read_csv(MER_CSV_DIR / "mer_output_files_inventory.csv")
    adrian_questions = read_csv(MER_CSV_DIR / "mer_adrian_questions.csv")

    counts = {row["agreement_category"]: row["plot_count"] for row in agreement}

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("Gayini MER methods and interpretation note")
    title_run.font.name = "Arial"
    title_run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string("20302C")
    doc.add_paragraph("Daily observed inundation footprint metrics as a complementary analysis")
    doc.add_paragraph("Generated 2026-06-26 | Codex Task 9 | D:/Github_repos/Gayini")

    doc.add_heading("1. Purpose of the MER analysis", level=1)
    doc.add_paragraph(
        "The MER / Flow-MER-inspired analysis was reviewed to determine whether daily observed inundation "
        "footprint metrics add useful context beyond the main Gayini annual occurrence and pre/post workflow. "
        "The current recommendation is to treat MER as supplementary decision support rather than a replacement "
        "for annual occurrence frequency."
    )

    doc.add_heading("2. Relationship to the main Gayini remote-sensing workflow", level=1)
    add_bullets(doc, [
        "Annual occurrence frequency remains the main inundation-change metric.",
        "Pre/post annual occurrence change remains the primary management-transition result.",
        "Historical background frequency and matched-year comparisons remain core interpretation supports.",
        "Gauge context remains hydrological support and is not part of the MER calculation.",
        "Ground-cover response interpretation remains separate and should not use MER as causal proof.",
    ])

    doc.add_heading("3. Input datasets", level=1)
    add_table(
        doc,
        ["File", "Type", "Role", "Date range", "Spatial unit", "Notes"],
        [[Path(r["file_path"]).name, r["file_type"], r["role"], r["date_range"], r["spatial_unit"], r["notes"]] for r in input_files],
        [1.1, 0.55, 1.35, 0.85, 0.85, 1.8],
    )

    doc.add_heading("4. Processing steps reviewed", level=1)
    add_bullets(doc, [
        "Reviewed the active driver scripts/06_mer/06_extract_MER_inundation_metrics.R.",
        "Reviewed the active implementation R/gayini_mer_inundation_functions.R.",
        "Reviewed the Task 4 consolidation handoff and Task 8 refreshed MER comparison assets.",
        "Built Task 9 tables and figure copies from existing outputs only.",
        "Did not rerun daily extraction, MER extraction, annual extraction or raster processing.",
    ])

    doc.add_heading("5. Metric definitions", level=1)
    add_table(
        doc,
        ["Metric", "Family", "Units", "Interpretation", "Caveat", "Use"],
        [[r["metric_name"], r["metric_family"], r["units"], r["interpretation"], r["main_caveat"], r["main_deck_appendix_defer"]] for r in metric_defs],
        [1.25, 1.0, 0.55, 1.5, 1.55, 0.9],
    )

    doc.add_heading("6. Annual occurrence versus MER-style metrics", level=1)
    doc.add_paragraph(
        "Annual occurrence frequency asks whether inundation was detected at least once in a valid water year. "
        "MER annual maximum observed area asks how much of the plot was observed wet at the largest detected event "
        "within the year. Neither metric is flood depth. Neither metric is full hydroperiod unless a future workflow "
        "explicitly calculates duration from sufficiently dense observations."
    )

    doc.add_heading("7. Agreement / disagreement result", level=1)
    doc.add_paragraph(
        f"The current MER versus annual occurrence comparison covers 66 plots: "
        f"{counts.get('Directions agree', '48')} directions agree, "
        f"{counts.get('Directions disagree / review', '12')} directions disagree / require review, and "
        f"{counts.get('One metric near no change', '6')} have one metric near no change."
    )
    add_table(
        doc,
        ["Category", "Plots", "Interpretation", "Review action"],
        [[r["agreement_category"], r["plot_count"], r["interpretation"], r["recommended_review_action"]] for r in agreement],
        [1.3, 0.45, 2.1, 2.0],
    )

    doc.add_heading("8. Outputs produced", level=1)
    add_table(
        doc,
        ["File", "Type", "Metric", "Use", "Location"],
        [[Path(r["file_path"]).name, r["file_type"], r["metric"], r["recommended_use"], r["deck_or_appendix"]] for r in output_files],
        [1.7, 0.55, 1.25, 1.6, 1.15],
    )

    doc.add_heading("9. Interpretation", level=1)
    doc.add_paragraph(
        "MER outputs currently support the view that daily observed footprint metrics add interpretive context. "
        "They are especially useful for understanding whether the largest observed wet footprints became larger or "
        "smaller and for identifying plots where MER and annual occurrence tell different stories. This strengthens "
        "review triage but does not alter the main annual occurrence framework."
    )
    add_figure(
        doc,
        MER_FIG_DIR / "mer_vs_annual_occurrence_main_deck_comparison.png",
        "MER versus annual occurrence comparison. Current spatial display is plot-based / centroid-based, not a continuous MER raster surface.",
    )

    doc.add_heading("10. Limitations", level=1)
    add_bullets(doc, [
        "The current MER comparison is plot-based / plot-centroid based because no true MER raster surface was found.",
        "Sensor observation support differs across years; Landsat and Sentinel-2 differ in cadence, pixel size and history.",
        "MER annual maximum observed area is not duration, hydroperiod or wet days.",
        "Annual occurrence frequency is not hydroperiod, duration or depth.",
        "Disagreement between metrics is a review flag, not automatically an error.",
        "MER outputs should not be used as causal evidence of management effects.",
    ])

    doc.add_heading("11. Recommended use in main review deck", level=1)
    doc.add_paragraph(
        "Use one MER comparison slide only if space allows. The recommended main-deck wording is "
        "'daily observed inundation footprint metrics'. The slide should emphasise complementarity: annual occurrence "
        "remains the headline result, while MER adds event-footprint context and review flags."
    )

    doc.add_heading("12. Recommended use in technical appendix", level=1)
    add_bullets(doc, [
        "Metric definitions and input/output inventory.",
        "Observation support by year/sensor.",
        "Detailed agreement/disagreement table.",
        "Selected plot-review flags.",
        "Deferred metric inventory for monthly/seasonal and sequence diagnostics.",
    ])

    doc.add_heading("13. Outstanding questions for Adrian", level=1)
    add_numbered(doc, [r["question"] for r in adrian_questions])

    doc.add_heading("14. File inventory", level=1)
    add_bullets(doc, [
        "Output/reports/MER/Gayini_MER_analysis_review_deck.pptx",
        "Output/reports/MER/Gayini_MER_methods_and_interpretation_note.docx",
        "Output/reports/MER/Gayini_MER_analysis_workbook.xlsx",
        "Output/csv/MER/mer_vs_annual_occurrence_agreement_summary.csv",
        "Output/csv/MER/mer_vs_annual_occurrence_plot_review_flags.csv",
        "Output/figures/review/MER/mer_vs_annual_occurrence_main_deck_comparison.png",
        "Output/figures/review/MER/mer_metric_summary_main_deck.png",
    ])

    doc.add_heading("Decision point on terminology", level=1)
    doc.add_paragraph(
        "Recommended wording: use 'daily observed inundation footprint metrics' in the main review deck, and "
        "'MER / Flow-MER-inspired inundation metrics' in the technical deck, Word note and workbook. Adrian should "
        "confirm the preferred label before the final public-facing deck is rebuilt."
    )

    doc.save(OUT_DOCX)
    print(f"Wrote: {OUT_DOCX}")


if __name__ == "__main__":
    main()
